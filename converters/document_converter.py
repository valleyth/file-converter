import platform
from pathlib import Path

import fitz
from docx import Document

from core.converter import BaseConverter


class DocumentConverter(BaseConverter):

    def supported_input(self) -> list[str]:
        return [".pdf", ".txt", ".docx"]

    def supported_output(self) -> list[str]:
        return [".pdf", ".txt", ".docx"]

    def convert(self, input_path: Path, output_path: Path, options: dict) -> None:
        input_ext = input_path.suffix.lower()
        output_ext = output_path.suffix.lower()

        if input_ext == ".pdf" and output_ext == ".txt":
            self._pdf_to_txt(input_path, output_path)
        elif input_ext == ".txt" and output_ext == ".pdf":
            self._txt_to_pdf(input_path, output_path)
        elif input_ext == ".docx" and output_ext == ".txt":
            self._docx_to_txt(input_path, output_path)
        elif input_ext == ".txt" and output_ext == ".docx":
            self._txt_to_docx(input_path, output_path)
        elif input_ext == ".docx" and output_ext == ".pdf":
            self._docx_to_pdf(input_path, output_path)
        else:
            raise ValueError(f"Неподдерживаемая конвертация: {input_ext} → {output_ext}")

    def _docx_to_txt(self, input_path: Path, output_path: Path) -> None:
        doc = Document(str(input_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        output_path.write_text(text, encoding="utf-8")

    def _txt_to_docx(self, input_path: Path, output_path: Path) -> None:
        text = input_path.read_text(encoding="utf-8")
        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(str(output_path))

    def _docx_to_pdf(self, input_path: Path, output_path: Path) -> None:
        if platform.system() == "Windows":
            self._docx_to_pdf_win(input_path, output_path)
        else:
            self._docx_to_pdf_linux(input_path, output_path)

    def _docx_to_pdf_win(self, input_path: Path, output_path: Path) -> None:
        import win32com.client

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(input_path))
        doc.SaveAs(str(output_path), FileFormat=17)
        doc.Close()
        word.Quit()

    def _docx_to_pdf_linux(self, input_path: Path, output_path: Path) -> None:
        import subprocess

        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir",
             str(output_path.parent), str(input_path)],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice ошибка: {result.stderr}")

    def _pdf_to_txt(self, input_path: Path, output_path: Path) -> None:
        doc = fitz.open(str(input_path))
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        output_path.write_text(text, encoding="utf-8")

    def _txt_to_pdf(self, input_path: Path, output_path: Path) -> None:
        from fpdf import FPDF

        text = input_path.read_text(encoding="utf-8")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 10, text)
        pdf.output(str(output_path))